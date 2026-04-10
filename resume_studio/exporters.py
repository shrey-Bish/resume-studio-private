from __future__ import annotations

import re
import zipfile
from html import escape
from io import BytesIO

from docx import Document


def markdown_bytes(value: str) -> bytes:
    return value.encode("utf-8")


def docx_bytes(title: str, body: str) -> bytes:
    document = Document()
    document.add_heading(title, level=1)

    for block in body.split("\n"):
        line = block.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def pdf_bytes(title: str, body: str) -> bytes:
    reportlab = _load_reportlab()

    buffer = BytesIO()
    document = reportlab["SimpleDocTemplate"](
        buffer,
        pagesize=reportlab["A4"],
        leftMargin=0.55 * reportlab["inch"],
        rightMargin=0.55 * reportlab["inch"],
        topMargin=0.55 * reportlab["inch"],
        bottomMargin=0.55 * reportlab["inch"],
    )
    styles = _pdf_styles(reportlab)
    story = [reportlab["Paragraph"](_escape_pdf(title), styles["title"]), reportlab["Spacer"](1, 10)]

    bullet_items: list[object] = []
    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line:
            if bullet_items:
                story.append(reportlab["ListFlowable"](bullet_items, bulletType="bullet", leftIndent=18))
                bullet_items = []
            story.append(reportlab["Spacer"](1, 6))
            continue

        if line.startswith("- "):
            bullet_items.append(
                reportlab["ListItem"](
                    reportlab["Paragraph"](_format_pdf_inline(line[2:].strip()), styles["body"])
                )
            )
            continue

        if bullet_items:
            story.append(reportlab["ListFlowable"](bullet_items, bulletType="bullet", leftIndent=18))
            bullet_items = []

        if line.startswith("# "):
            story.append(reportlab["Paragraph"](_escape_pdf(line[2:].strip()), styles["h1"]))
        elif line.startswith("## "):
            story.append(reportlab["Paragraph"](_escape_pdf(line[3:].strip()), styles["h2"]))
        elif line.startswith("### "):
            story.append(reportlab["Paragraph"](_escape_pdf(line[4:].strip()), styles["h3"]))
        else:
            story.append(reportlab["Paragraph"](_format_pdf_inline(line), styles["body"]))
        story.append(reportlab["Spacer"](1, 4))

    if bullet_items:
        story.append(reportlab["ListFlowable"](bullet_items, bulletType="bullet", leftIndent=18))

    document.build(story)
    buffer.seek(0)
    return buffer.read()


def pdf_export_supported() -> bool:
    try:
        _load_reportlab()
        return True
    except ModuleNotFoundError:
        return False


def _format_inline(value: str) -> str:
    escaped = escape(value)
    escaped = escaped.replace("**", "__BOLD__")
    escaped = escaped.replace("*", "__ITALIC__")

    parts = escaped.split("__BOLD__")
    for index in range(1, len(parts), 2):
        parts[index] = f"<strong>{parts[index]}</strong>"
    escaped = "".join(parts)

    parts = escaped.split("__ITALIC__")
    for index in range(1, len(parts), 2):
        parts[index] = f"<em>{parts[index]}</em>"
    return "".join(parts)


def zip_bytes(files: list[tuple[str, bytes]]) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, payload in files:
            archive.writestr(name, payload)
    buffer.seek(0)
    return buffer.read()


def _pdf_styles(reportlab: dict[str, object]) -> dict[str, object]:
    base = reportlab["getSampleStyleSheet"]()
    return {
        "title": reportlab["ParagraphStyle"](
            "Title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            textColor=reportlab["HexColor"]("#0f766e"),
            spaceAfter=8,
        ),
        "h1": reportlab["ParagraphStyle"](
            "H1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=reportlab["HexColor"]("#172033"),
            spaceBefore=8,
            spaceAfter=6,
        ),
        "h2": reportlab["ParagraphStyle"](
            "H2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=16,
            textColor=reportlab["HexColor"]("#172033"),
            spaceBefore=8,
            spaceAfter=4,
        ),
        "h3": reportlab["ParagraphStyle"](
            "H3",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=14,
            textColor=reportlab["HexColor"]("#172033"),
            spaceBefore=6,
            spaceAfter=3,
        ),
        "body": reportlab["ParagraphStyle"](
            "Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            textColor=reportlab["HexColor"]("#172033"),
        ),
    }


def _format_pdf_inline(value: str) -> str:
    text = _escape_pdf(value)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    return text


def _escape_pdf(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _load_reportlab() -> dict[str, object]:
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    return {
        "HexColor": HexColor,
        "A4": A4,
        "ParagraphStyle": ParagraphStyle,
        "getSampleStyleSheet": getSampleStyleSheet,
        "inch": inch,
        "ListFlowable": ListFlowable,
        "ListItem": ListItem,
        "Paragraph": Paragraph,
        "SimpleDocTemplate": SimpleDocTemplate,
        "Spacer": Spacer,
    }
